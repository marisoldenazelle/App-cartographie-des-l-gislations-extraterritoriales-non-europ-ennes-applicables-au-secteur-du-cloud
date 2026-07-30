# -*- coding: utf-8 -*-
"""
Cartographie interactive des législations extraterritoriales non européennes
applicables au secteur du cloud — DGE.
"""
import json
import re
from pathlib import Path
import io

import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
import streamlit as st
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ----------------------------------------------------------------------------
# Configuration & données
# ----------------------------------------------------------------------------
st.set_page_config(
    page_title="Cartographie extraterritorialité cloud — DGE",
    page_icon="🌐",
    layout="wide",
    initial_sidebar_state="expanded",
)

DATA_PATH = Path(__file__).parent / "data" / "note_data.json"


@st.cache_data
def load_data():
    with open(DATA_PATH, encoding="utf-8") as f:
        return json.load(f)


DATA = load_data()

RISK_COLORS = {
    "Élevé": "#C0392B",
    "Moyen": "#E08E0B",
    "Faible à moyen": "#D4AC0D",
    "Faible": "#1E8449",
}
RISK_ORDER = ["Faible", "Faible à moyen", "Moyen", "Élevé"]

# --- Catégorisation de la nature de la contrainte (2e dimension d'analyse) ---
# Mappage explicite par texte (source de vérité), avec repli par mots-clés si
# de nouvelles lignes sont ajoutées au JSON sans être répertoriées ici.
CATEGORIE_PAR_TEXTE = {
    "CLOUD Act": "Accès / communication de données",
    "FISA — art. 702": "Surveillance / interception",
    "Executive Order 12333": "Surveillance / interception",
    "Loi sur le renseignement national, art. 7": "Coopération / assistance",
    "Loi sur la sécurité des données (DSL)": "Conformité déclarative / sanctions",
    "PIPL": "Conformité déclarative / sanctions",
    "Investigatory Powers Act — TCN": "Backdoor / capacité technique",
    "Loi n° 242-FZ / 23-FZ": "Localisation des données",
    "TOLA (TAR/TAN/TCN)": "Backdoor / capacité technique",
    "Digital Personal Data Protection Act": "Conformité déclarative / sanctions",
    "Loi ISA / mémorandum 2023": "Surveillance / interception",
    "Personal Information Protection Act (PIPA)": "Conformité déclarative / sanctions",
    "Négociation accord exécutif CLOUD Act": "Accès / communication de données",
}
CATEGORIE_KEYWORD_FALLBACK = [
    (r"chiffrement|capacité technique|backdoor|porte dérobée", "Backdoor / capacité technique"),
    (r"localis", "Localisation des données"),
    (r"communication de données|demande judiciaire|accès réciproque", "Accès / communication de données"),
    (r"surveillance|renseignement.*collecte|interception|recherche à distance", "Surveillance / interception"),
    (r"coopération|assistance", "Coopération / assistance"),
    (r"conformité|sanctions", "Conformité déclarative / sanctions"),
]
CATEGORIE_COLORS = {
    "Accès / communication de données": "#2E86AB",
    "Surveillance / interception": "#A23B72",
    "Coopération / assistance": "#6A4C93",
    "Backdoor / capacité technique": "#C0392B",
    "Localisation des données": "#1B998B",
    "Conformité déclarative / sanctions": "#E08E0B",
}


def categorize_obligation(row: pd.Series) -> str:
    """Détermine la nature de la contrainte imposée par un texte (2e dimension,
    croisée avec le niveau de risque dans la heatmap)."""
    cat = CATEGORIE_PAR_TEXTE.get(row["texte"])
    if cat:
        return cat
    haystack = f"{row['obligation']} {row['fondement']}".lower()
    for pattern, fallback_cat in CATEGORIE_KEYWORD_FALLBACK:
        if re.search(pattern, haystack):
            return fallback_cat
    return "Autre"


def extract_years(annee_str: str):
    """Extrait (année d'origine, année la plus récente mentionnée) d'un champ
    'annee' pouvant contenir un ou plusieurs millésimes (ex. '2015 / 2025')."""
    years = [int(y) for y in re.findall(r"(?:19|20)\d{2}", str(annee_str))]
    if not years:
        return None, None
    return min(years), max(years)


@st.cache_data
def build_tableau_df(tableau_comparatif: list) -> pd.DataFrame:
    df = pd.DataFrame(tableau_comparatif)
    df["categorie_obligation"] = df.apply(categorize_obligation, axis=1)
    years = df["annee"].apply(lambda a: pd.Series(extract_years(a)))
    df["annee_origine"], df["annee_recente"] = years[0], years[1]
    df["niveau_risque"] = pd.Categorical(df["niveau_risque"], categories=RISK_ORDER, ordered=True)
    return df

# ----------------------------------------------------------------------------
# Style
# ----------------------------------------------------------------------------
st.markdown(
    """
    <style>
    .main .block-container {padding-top: 2rem; max-width: 1200px;}
    h1, h2, h3 {font-family: 'Georgia', serif;}
    .note-header {
        background: linear-gradient(135deg, #0B3D66 0%, #14567F 100%);
        color: white; padding: 1.6rem 2rem; border-radius: 10px; margin-bottom: 1.5rem;
    }
    .note-header h1 {color: white; font-size: 1.5rem; margin-bottom: 0.3rem;}
    .note-header p {color: #D6E4F0; font-size: 0.95rem; margin: 0;}
    .country-card {
        border: 1px solid #DDD; border-left: 5px solid #14567F;
        border-radius: 6px; padding: 1rem 1.2rem; margin-bottom: 1rem; background: #FAFBFC;
    }
    .risk-badge {
        display:inline-block; padding: 2px 10px; border-radius: 12px;
        color: white; font-size: 0.78rem; font-weight: 600;
    }
    .source-link {font-size: 0.85rem;}
    mark {background-color: #FFEB99; padding: 0 2px;}
    </style>
    """,
    unsafe_allow_html=True,
)

# ----------------------------------------------------------------------------
# Sidebar — navigation
# ----------------------------------------------------------------------------
SECTIONS = [
    "🏠 Synthèse",
    "🔍 Recherche plein texte",
    "📖 Méthodologie",
    "🗺️ Panorama par juridiction",
    "📊 Tableau comparatif interactif",
    "⚠️ Analyse transversale des risques",
    "🇪🇺 Réponses réglementaires FR/UE",
    "📚 Sources",
    "📄 Export Word",
]

st.sidebar.title("🌐 Navigation")
section = st.sidebar.radio("Aller à la section :", SECTIONS, label_visibility="collapsed")
st.sidebar.markdown("---")
st.sidebar.caption(
    f"**{DATA['meta']['produced_by']}** \n"
    f"Note d'analyse — {DATA['meta']['date_note']}"
)
st.sidebar.caption(
    "Cette application est une mise en forme interactive de la note d'analyse. "
    "Elle ne constitue pas une consultation juridique."
)

# ----------------------------------------------------------------------------
# En-tête commun
# ----------------------------------------------------------------------------
st.markdown(
    f"""
    <div class="note-header">
        <h1>{DATA['meta']['title']}</h1>
        <p>{DATA['meta']['subtitle']}</p>
    </div>
    """,
    unsafe_allow_html=True,
)


def risk_badge(level: str) -> str:
    color = RISK_COLORS.get(level, "#888")
    return f'<span class="risk-badge" style="background:{color};">{level}</span>'


def source_link(url, label):
    if url:
        return f'<a class="source-link" href="{url}" target="_blank">🔗 {label}</a>'
    return '<span class="source-link" style="color:#999;">Source primaire non disponible en ligne — voir Annexe Sources</span>'


# ----------------------------------------------------------------------------
# SECTION — Synthèse
# ----------------------------------------------------------------------------
if section == "🏠 Synthèse":
    for para in DATA["synthese"]["abstract"]:
        st.write(para)

    st.markdown("### Les six régimes juridiques majeurs identifiés")
    df_regimes = pd.DataFrame(DATA["synthese"]["regimes_cles"])
    df_regimes.columns = ["Pays", "Texte de référence", "Année"]
    st.dataframe(df_regimes, use_container_width=True, hide_index=True)

    st.info(
        "💡 Utilisez la navigation à gauche pour explorer le panorama détaillé par "
        "juridiction, filtrer le tableau comparatif, ou rechercher un mot-clé dans "
        "l'ensemble de la note.",
        icon="💡",
    )

# ----------------------------------------------------------------------------
# SECTION — Recherche plein texte
# ----------------------------------------------------------------------------
elif section == "🔍 Recherche plein texte":
    st.subheader("Recherche dans l'ensemble de la note")
    query = st.text_input(
        "Mot-clé ou expression (ex. « FISA », « chiffrement », « SecNumCloud »)",
        placeholder="Tapez votre recherche…",
    )

    def collect_searchable_blocks():
        blocks = []
        for p in DATA["synthese"]["abstract"]:
            blocks.append(("Synthèse", "Synthèse", p))
        for key, label in [
            ("objet", "Méthodologie — Objet"),
            ("definition", "Méthodologie — Définition"),
            ("perimetre", "Méthodologie — Périmètre"),
            ("sources_limites", "Méthodologie — Sources et limites"),
        ]:
            blocks.append(("Méthodologie", label, DATA["methodologie"][key]))
        for f in DATA["methodologie"]["fondements"]:
            blocks.append(("Méthodologie", f["titre"], f["description"]))
        for pays in DATA["pays"]:
            if pays["intro"]:
                blocks.append((pays["nom"], f"{pays['nom']} — Introduction", pays["intro"]))
            for texte in pays["textes"]:
                blocks.append((pays["nom"], f"{pays['nom']} — {texte['nom']}", texte["description"]))
                for extrait in texte.get("extraits", []):
                    blocks.append((
                        pays["nom"],
                        f"{pays['nom']} — {texte['nom']} — {extrait['reference']}",
                        extrait["resume"],
                    ))
            if pays.get("note_complementaire"):
                blocks.append((pays["nom"], f"{pays['nom']} — Note complémentaire", pays["note_complementaire"]))
        for r in DATA["risques_transversaux"]:
            blocks.append(("Analyse transversale", r["titre"], r["texte"]))
        for r in DATA["reponses_reglementaires"]:
            blocks.append(("Réponses réglementaires", r["titre"], r["texte"]))
        return blocks

    blocks = collect_searchable_blocks()

    if query and query.strip():
        pattern = re.compile(re.escape(query.strip()), re.IGNORECASE)
        results = [(sec, title, text) for sec, title, text in blocks if pattern.search(text)]
        st.caption(f"{len(results)} passage(s) trouvé(s) pour « {query} »")
        for sec, title, text in results:
            highlighted = pattern.sub(lambda m: f"<mark>{m.group(0)}</mark>", text)
            with st.container():
                st.markdown(f"**{title}** &nbsp;·&nbsp; *{sec}*")
                st.markdown(f"<div style='margin-bottom:1.2rem;'>{highlighted}</div>", unsafe_allow_html=True)
    else:
        st.caption("Saisissez un terme pour lancer la recherche dans les 9 sections de la note.")

# ----------------------------------------------------------------------------
# SECTION — Méthodologie
# ----------------------------------------------------------------------------
elif section == "📖 Méthodologie":
    st.subheader("1. Objet, périmètre et méthodologie")
    st.markdown("**Objet de la note**")
    st.write(DATA["methodologie"]["objet"])

    st.markdown("**1.1 Définition retenue de l'extraterritorialité**")
    st.write(DATA["methodologie"]["definition"])

    st.markdown("Trois fondements de compétence extraterritoriale sont observés :")
    cols = st.columns(3)
    for col, f in zip(cols, DATA["methodologie"]["fondements"]):
        with col:
            st.markdown(f"**{f['titre']}**")
            st.caption(f["description"])

    st.markdown("**1.2 Périmètre géographique**")
    st.write(DATA["methodologie"]["perimetre"])

    st.markdown("**1.3 Sources et limites de l'analyse**")
    st.write(DATA["methodologie"]["sources_limites"])

# ----------------------------------------------------------------------------
# SECTION — Panorama par juridiction
# ----------------------------------------------------------------------------
elif section == "🗺️ Panorama par juridiction":
    st.subheader("2. Panorama des régimes juridiques extraterritoriaux (hors UE)")

    noms = [f"{p['drapeau']} {p['nom']}" for p in DATA["pays"]]
    tabs = st.tabs(noms)

    for tab, pays in zip(tabs, DATA["pays"]):
        with tab:
            if pays["intro"]:
                st.write(pays["intro"])
            for texte in pays["textes"]:
                with st.expander(f"📄 {texte['nom']} ({texte['annee']})", expanded=len(pays["textes"]) == 1):
                    st.write(texte["description"])
                    extraits = texte.get("extraits") or []
                    if extraits:
                        st.markdown("**Extraits clés :**")
                        for ex in extraits:
                            st.markdown(f"- **{ex['reference']}** — {ex['resume']}")
                    if texte.get("url"):
                        st.markdown(
                            source_link(texte["url"], texte.get("source_label") or "Accéder au texte"),
                            unsafe_allow_html=True,
                        )
                    else:
                        st.caption("Source primaire non disponible en ligne — voir section Sources.")
            if pays.get("note_complementaire"):
                st.info(pays["note_complementaire"])
            st.markdown("---")

    st.markdown("**2.7 Autres juridictions à surveiller**")
    st.write(DATA["autres_juridictions"])

# ----------------------------------------------------------------------------
# SECTION — Tableau comparatif interactif
# ----------------------------------------------------------------------------
elif section == "📊 Tableau comparatif interactif":
    st.subheader("3. Tableau comparatif de synthèse")
    st.caption(DATA["tableau_note"])

    df_all = build_tableau_df(DATA["tableau_comparatif"])
    id_to_pays = {p["id"]: p for p in DATA["pays"]}

    def find_source(row):
        pays_obj = id_to_pays.get(row["pays_id"])
        if not pays_obj:
            return None, None
        for t in pays_obj["textes"]:
            if t["nom"].split(" (")[0][:20].lower() in row["texte"].lower() or row["texte"].split(" —")[0][:15].lower() in t["nom"].lower():
                return t.get("url"), t.get("source_label")
        return None, None

    df_all["url"], df_all["source_label"] = zip(*df_all.apply(find_source, axis=1))

    # --- Filtres communs à toutes les vues ---
    fc1, fc2, fc3, fc4 = st.columns([2, 2, 2, 3])
    with fc1:
        juridictions_sel = st.multiselect(
            "Juridiction", options=sorted(df_all["juridiction"].unique()),
            placeholder="Toutes",
        )
    with fc2:
        risques_sel = st.multiselect(
            "Niveau de risque", options=RISK_ORDER, placeholder="Tous",
        )
    with fc3:
        categories_sel = st.multiselect(
            "Nature de la contrainte", options=sorted(df_all["categorie_obligation"].unique()),
            placeholder="Toutes",
        )
    with fc4:
        recherche = st.text_input("Recherche libre", placeholder="ex. chiffrement, backdoor, PIPL…")

    df = df_all.copy()
    if juridictions_sel:
        df = df[df["juridiction"].isin(juridictions_sel)]
    if risques_sel:
        df = df[df["niveau_risque"].isin(risques_sel)]
    if categories_sel:
        df = df[df["categorie_obligation"].isin(categories_sel)]
    if recherche.strip():
        pat = re.escape(recherche.strip())
        mask = (
            df["texte"].str.contains(pat, case=False, na=False)
            | df["fondement"].str.contains(pat, case=False, na=False)
            | df["obligation"].str.contains(pat, case=False, na=False)
        )
        df = df[mask]

    if df.empty:
        st.warning("Aucun texte ne correspond à ces filtres.")
        st.stop()

    st.markdown(f"**{len(df)} texte(s) affiché(s) sur {len(df_all)}**")

    # --- KPI ---
    k1, k2, k3, k4 = st.columns(4)
    k1.metric("Textes recensés", len(df))
    k2.metric("Juridictions couvertes", df["juridiction"].nunique())
    k3.metric("Risque élevé", int((df["niveau_risque"] == "Élevé").sum()))
    k4.metric("Natures de contrainte", df["categorie_obligation"].nunique())

    tab_heat, tab_time, tab_radar, tab_table = st.tabs(
        ["🔥 Heatmap risque × obligation", "🕰️ Frise chronologique", "🕸️ Comparatif par juridiction", "📋 Tableau détaillé"]
    )

    # ============================ HEATMAP ============================
    with tab_heat:
        st.markdown(
            "Chaque cellule croise une **juridiction** et une **nature de contrainte** ; "
            "la couleur indique le **niveau de risque maximal** observé pour cette combinaison. "
            "Les cases vides signifient qu'aucun texte recensé dans cette juridiction ne relève de cette catégorie."
        )
        pivot = df.pivot_table(
            index="juridiction", columns="categorie_obligation", values="score_risque", aggfunc="max"
        )
        pivot = pivot.sort_index(axis=1)
        # Trie les juridictions par sévérité décroissante pour une lecture immédiate
        pivot = pivot.loc[pivot.max(axis=1).sort_values(ascending=False).index]

        # Texte affiché dans chaque cellule = niveau de risque en toutes lettres
        score_to_label = {3: "Élevé", 2: "Moyen", 1.5: "Faible à moyen", 1: "Faible"}
        label_fn = lambda v: score_to_label.get(v, "") if pd.notna(v) else ""
        try:
            text_matrix = pivot.map(label_fn)  # pandas >= 2.1
        except AttributeError:
            text_matrix = pivot.applymap(label_fn)  # pandas < 2.1

        fig_heat = go.Figure(
            data=go.Heatmap(
                z=pivot.values,
                x=pivot.columns,
                y=pivot.index,
                text=text_matrix.values,
                texttemplate="%{text}",
                textfont={"size": 11},
                colorscale=[[0, "#1E8449"], [0.35, "#D4AC0D"], [0.65, "#E08E0B"], [1, "#C0392B"]],
                zmin=1, zmax=3,
                hoverongaps=False,
                hovertemplate="<b>%{y}</b><br>%{x}<br>Risque max : %{text}<extra></extra>",
                colorbar=dict(title="Risque", tickvals=[1, 1.5, 2, 3], ticktext=["Faible", "Faible/moyen", "Moyen", "Élevé"]),
            )
        )
        fig_heat.update_layout(
            height=max(360, 60 + 34 * pivot.shape[0]),
            margin=dict(l=10, r=10, t=10, b=10),
            xaxis=dict(side="top", tickangle=-20),
        )
        st.plotly_chart(fig_heat, use_container_width=True)
        st.caption(
            "Lecture : les juridictions dont la ligne est chargée sur plusieurs colonnes cumulent plusieurs "
            "leviers d'accès (ex. surveillance **et** capacité technique). Une colonne concentrée sur peu de "
            "juridictions signale un mode d'action moins répandu mais pas nécessairement moins risqué."
        )

    # ============================ FRISE ============================
    with tab_time:
        st.markdown(
            "Positionnement des textes selon leur **année d'origine** ; la taille et la couleur du marqueur "
            "reflètent le niveau de risque. Survolez un point pour voir l'année de renforcement la plus récente, "
            "le cas échéant."
        )
        df_time = df.dropna(subset=["annee_origine"]).copy()
        df_time["renforce"] = df_time["annee_recente"] > df_time["annee_origine"]
        df_time["hover"] = df_time.apply(
            lambda r: (
                f"{r['texte']}<br>Origine : {int(r['annee_origine'])}"
                + (f" · dernière évolution : {int(r['annee_recente'])}" if r["renforce"] else "")
                + f"<br>{r['categorie_obligation']}"
            ),
            axis=1,
        )
        fig_time = px.scatter(
            df_time, x="annee_origine", y="juridiction", color="niveau_risque",
            color_discrete_map=RISK_COLORS, size="score_risque", size_max=22,
            category_orders={"niveau_risque": RISK_ORDER},
            custom_data=["hover"],
        )
        fig_time.update_traces(hovertemplate="%{customdata[0]}<extra></extra>")
        # Marqueur distinct pour les textes récemment renforcés
        reinforced = df_time[df_time["renforce"]]
        if not reinforced.empty:
            fig_time.add_trace(go.Scatter(
                x=reinforced["annee_recente"], y=reinforced["juridiction"], mode="markers",
                marker=dict(symbol="star", size=13, color="#0B3D66", line=dict(width=1, color="white")),
                name="Renforcement récent", hovertext=reinforced["texte"], hoverinfo="text",
            ))
        fig_time.update_layout(
            height=max(360, 50 + 40 * df_time["juridiction"].nunique()),
            xaxis_title="Année", yaxis_title=None,
            margin=dict(l=10, r=10, t=30, b=10),
            legend_title_text="Niveau de risque",
        )
        st.plotly_chart(fig_time, use_container_width=True)
        st.caption("⭐ = texte modifié ou renforcé après son année d'entrée en vigueur d'origine.")

    # ============================ RADAR ============================
    with tab_radar:
        st.caption(
            "⚠️ Ces axes sont des indicateurs **dérivés** du tableau comparatif (score de risque moyen/maximal, "
            "nombre de textes recensés) — une lecture de synthèse ajoutée pour faciliter la comparaison visuelle, "
            "et non une notation qualitative supplémentaire issue de la note d'origine."
        )
        agg = df.groupby("juridiction").agg(
            score_moyen=("score_risque", "mean"),
            score_max=("score_risque", "max"),
            nb_textes=("texte", "count"),
        ).reset_index()
        max_nb = agg["nb_textes"].max()
        agg["nb_textes_norm"] = agg["nb_textes"] / max_nb * 3 if max_nb else 0

        default_sel = agg.sort_values("score_max", ascending=False)["juridiction"].head(4).tolist()
        radar_sel = st.multiselect(
            "Juridictions à comparer (4 maximum recommandé pour la lisibilité)",
            options=sorted(agg["juridiction"]), default=default_sel,
        )
        if radar_sel:
            fig_radar = go.Figure()
            for j in radar_sel:
                r = agg[agg["juridiction"] == j].iloc[0]
                fig_radar.add_trace(go.Scatterpolar(
                    r=[r["score_moyen"], r["score_max"], r["nb_textes_norm"], r["score_moyen"]],
                    theta=["Score de risque moyen", "Score de risque maximal", "Nb. de textes (indexé /3)", "Score de risque moyen"],
                    fill="toself", name=j, opacity=0.75,
                ))
            fig_radar.update_layout(
                polar=dict(radialaxis=dict(visible=True, range=[0, 3])),
                height=480, margin=dict(l=40, r=40, t=30, b=30),
                legend=dict(orientation="h", yanchor="bottom", y=-0.15),
            )
            st.plotly_chart(fig_radar, use_container_width=True)
        else:
            st.info("Sélectionnez au moins une juridiction pour afficher le radar.")

        with st.expander("Détail des indicateurs agrégés par juridiction"):
            st.dataframe(
                agg.sort_values("score_max", ascending=False).rename(columns={
                    "juridiction": "Juridiction", "score_moyen": "Score moyen", "score_max": "Score max",
                    "nb_textes": "Nb. de textes", "nb_textes_norm": "Nb. de textes (indexé /3)",
                }),
                use_container_width=True, hide_index=True,
            )

    # ============================ TABLEAU ============================
    with tab_table:
        sort_col = st.selectbox(
            "Trier par :", ["Niveau de risque (décroissant)", "Juridiction (A→Z)", "Année (récente d'abord)"],
        )
        df_sorted = df.copy()
        if sort_col == "Niveau de risque (décroissant)":
            df_sorted = df_sorted.sort_values("score_risque", ascending=False)
        elif sort_col == "Juridiction (A→Z)":
            df_sorted = df_sorted.sort_values("juridiction")
        else:
            df_sorted = df_sorted.sort_values("annee_recente", ascending=False)

        display_df = df_sorted[[
            "juridiction", "texte", "annee", "categorie_obligation", "niveau_risque",
            "fondement", "obligation", "url",
        ]].rename(columns={
            "juridiction": "Juridiction", "texte": "Texte", "annee": "Année",
            "categorie_obligation": "Nature de la contrainte", "niveau_risque": "Risque",
            "fondement": "Fondement", "obligation": "Obligation", "url": "Source",
        })

        def highlight_risk(val):
            return f"background-color: {RISK_COLORS.get(val, '#FFF')}22; font-weight:600;" if val in RISK_COLORS else ""

        try:
            styled = display_df.style.map(highlight_risk, subset=["Risque"])  # pandas >= 2.1
        except AttributeError:
            styled = display_df.style.applymap(highlight_risk, subset=["Risque"])  # pandas < 2.1

        st.dataframe(
            styled,
            use_container_width=True, hide_index=True,
            column_config={
                "Source": st.column_config.LinkColumn("Source", display_text="🔗 Texte source"),
                "Fondement": st.column_config.TextColumn("Fondement", width="medium"),
                "Obligation": st.column_config.TextColumn("Obligation", width="medium"),
            },
            height=min(600, 70 + 36 * len(display_df)),
        )

        csv = display_df.drop(columns=["Source"]).to_csv(index=False).encode("utf-8-sig")
        st.download_button(
            "⬇️ Exporter le tableau filtré (CSV)", data=csv,
            file_name="tableau_comparatif_extraterritorialite_cloud.csv", mime="text/csv",
        )

# ----------------------------------------------------------------------------
# SECTION — Analyse transversale
# ----------------------------------------------------------------------------
elif section == "⚠️ Analyse transversale des risques":
    st.subheader("4. Analyse transversale des risques pour le secteur du cloud")
    for i, r in enumerate(DATA["risques_transversaux"], 1):
        with st.expander(f"4.{i} {r['titre']}", expanded=True):
            st.write(r["texte"])

# ----------------------------------------------------------------------------
# SECTION — Réponses réglementaires
# ----------------------------------------------------------------------------
elif section == "🇪🇺 Réponses réglementaires FR/UE":
    st.subheader("5. Réponses réglementaires françaises et européennes")
    for i, r in enumerate(DATA["reponses_reglementaires"], 1):
        st.markdown(f"**5.{i} {r['titre']}**")
        st.write(r["texte"])
        if r.get("url"):
            st.markdown(source_link(r["url"], r.get("source_label") or "Accéder à la source"), unsafe_allow_html=True)
        st.markdown("")

# ----------------------------------------------------------------------------
# SECTION — Sources
# ----------------------------------------------------------------------------
elif section == "📚 Sources":
    st.subheader("Annexe — Sources")
    st.caption(DATA["sources_note"])
    for s in DATA["sources"]:
        if s.get("url"):
            st.markdown(f"- [{s['nom']}]({s['url']})")
        else:
            st.markdown(f"- {s['nom']}")

# ----------------------------------------------------------------------------
# SECTION — Export Word ajustable
# ----------------------------------------------------------------------------
elif section == "📄 Export Word":
    st.subheader("Générer une synthèse Word sur mesure")
    st.caption(
        "Sélectionnez les éléments à inclure, puis générez un document .docx prêt à partager "
        "ou à annoter."
    )

    with st.form("export_form"):
        st.markdown("**Sections à inclure**")
        c1, c2 = st.columns(2)
        with c1:
            inc_synthese = st.checkbox("Synthèse", value=True)
            inc_methodo = st.checkbox("Méthodologie", value=False)
            inc_panorama = st.checkbox("Panorama par juridiction", value=True)
            inc_tableau = st.checkbox("Tableau comparatif", value=True)
        with c2:
            inc_risques = st.checkbox("Analyse transversale des risques", value=True)
            inc_reponses = st.checkbox("Réponses réglementaires FR/UE", value=True)
            inc_sources = st.checkbox("Sources", value=False)

        pays_options = [p["nom"] for p in DATA["pays"]]
        pays_choisis = st.multiselect(
            "Juridictions à inclure dans le panorama (si sélectionné ci-dessus)",
            options=pays_options, default=pays_options,
        )

        submitted = st.form_submit_button("📄 Générer le document Word")

    def build_docx():
        doc = Document()
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(10.5)

        title = doc.add_heading(DATA["meta"]["title"], level=0)
        subtitle_p = doc.add_paragraph(DATA["meta"]["subtitle"])
        subtitle_p.runs[0].italic = True
        meta_p = doc.add_paragraph(f"{DATA['meta']['produced_by']} — {DATA['meta']['date_note']}")
        meta_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        if inc_synthese:
            doc.add_heading("Synthèse", level=1)
            for para in DATA["synthese"]["abstract"]:
                doc.add_paragraph(para)
            doc.add_heading("Les six régimes juridiques majeurs", level=2)
            table = doc.add_table(rows=1, cols=3)
            table.style = "Light Grid Accent 1"
            hdr = table.rows[0].cells
            hdr[0].text, hdr[1].text, hdr[2].text = "Pays", "Texte de référence", "Année"
            for r in DATA["synthese"]["regimes_cles"]:
                row = table.add_row().cells
                row[0].text, row[1].text, row[2].text = r["pays"], r["texte"], r["annee"]

        if inc_methodo:
            doc.add_heading("Méthodologie", level=1)
            doc.add_paragraph(DATA["methodologie"]["objet"])
            doc.add_heading("Définition retenue", level=2)
            doc.add_paragraph(DATA["methodologie"]["definition"])
            doc.add_heading("Fondements de compétence extraterritoriale", level=2)
            for f in DATA["methodologie"]["fondements"]:
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(f["titre"] + " — ").bold = True
                p.add_run(f["description"])
            doc.add_heading("Périmètre géographique", level=2)
            doc.add_paragraph(DATA["methodologie"]["perimetre"])
            doc.add_heading("Sources et limites", level=2)
            doc.add_paragraph(DATA["methodologie"]["sources_limites"])

        if inc_panorama:
            doc.add_heading("Panorama par juridiction", level=1)
            for pays in DATA["pays"]:
                if pays["nom"] not in pays_choisis:
                    continue
                doc.add_heading(f"{pays['drapeau']} {pays['nom']}", level=2)
                if pays["intro"]:
                    doc.add_paragraph(pays["intro"])
                for texte in pays["textes"]:
                    doc.add_heading(f"{texte['nom']} ({texte['annee']})", level=3)
                    doc.add_paragraph(texte["description"])
                    for ex in texte.get("extraits") or []:
                        p = doc.add_paragraph(style="List Bullet")
                        p.add_run(ex["reference"] + " — ").bold = True
                        p.add_run(ex["resume"])
                    if texte.get("url"):
                        doc.add_paragraph(f"Source : {texte.get('source_label') or texte['url']} ({texte['url']})")
                if pays.get("note_complementaire"):
                    note_p = doc.add_paragraph(pays["note_complementaire"])
                    note_p.runs[0].italic = True

        if inc_tableau:
            doc.add_heading("Tableau comparatif de synthèse", level=1)
            doc.add_paragraph(DATA["tableau_note"])
            cols = ["Juridiction", "Texte", "Année", "Fondement", "Obligation", "Niveau de risque"]
            table = doc.add_table(rows=1, cols=len(cols))
            table.style = "Light Grid Accent 1"
            for i, c in enumerate(cols):
                table.rows[0].cells[i].text = c
            for row in DATA["tableau_comparatif"]:
                cells = table.add_row().cells
                cells[0].text = row["juridiction"]
                cells[1].text = row["texte"]
                cells[2].text = row["annee"]
                cells[3].text = row["fondement"]
                cells[4].text = row["obligation"]
                cells[5].text = row["niveau_risque"]

        if inc_risques:
            doc.add_heading("Analyse transversale des risques", level=1)
            for r in DATA["risques_transversaux"]:
                doc.add_heading(r["titre"], level=2)
                doc.add_paragraph(r["texte"])

        if inc_reponses:
            doc.add_heading("Réponses réglementaires françaises et européennes", level=1)
            for r in DATA["reponses_reglementaires"]:
                doc.add_heading(r["titre"], level=2)
                doc.add_paragraph(r["texte"])
                if r.get("url"):
                    doc.add_paragraph(f"Source : {r.get('source_label') or r['url']} ({r['url']})")

        if inc_sources:
            doc.add_heading("Sources", level=1)
            doc.add_paragraph(DATA["sources_note"])
            for s in DATA["sources"]:
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(s["nom"] + (f" ({s['url']})" if s.get("url") else ""))

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf

    if submitted:
        if not any([inc_synthese, inc_methodo, inc_panorama, inc_tableau, inc_risques, inc_reponses, inc_sources]):
            st.warning("Sélectionnez au moins une section avant de générer le document.")
        else:
            buf = build_docx()
            st.success("Document généré ✅")
            st.download_button(
                label="⬇️ Télécharger le document Word",
                data=buf,
                file_name="synthese_extraterritorialite_cloud.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
